from data import TABLE_NAME, USER_KEY
from airtable import Airtable
from docx2pdf import convert
from docx import Document
import os


FOLDER_NAME = TABLE_NAME + ' Submissions'


def get_airtable_records(base_key, table_name):
    airtable = Airtable(
        base_key, table_name, USER_KEY
    )
    return airtable.get_all(sort = 'Division')


def create_empty_folders(divisions):
    if not FOLDER_NAME in os.listdir():
        os.mkdir(FOLDER_NAME)
        print('Created base folder.')
    else:
        print('Base folder exists.')
    os.chdir(FOLDER_NAME)
    
    delete_count = 0
    for division in divisions:
        if division not in os.listdir():
            os.mkdir(division)
        else:
            for file in os.listdir(division):
                os.remove(division + '/' + file)
                delete_count += 1
    print('Deleted previous files:', delete_count, '\n')


def create_pdf(test):
    try:
        document = Document()
        file_name = test['fields']['Division'] + '/' + test['fields']['Name'] + '.docx'
        for i in test['fields']:
            document.add_heading(str(i)).bold = True
            document.add_paragraph(test['fields'][i])
        document.add_paragraph(test['createdTime'])
        if test['fields']['Division'] in os.listdir():
            document.save(file_name)
            convert(file_name)
            os.remove(file_name)
        return True

    except Exception as e:
        print(e)
        return False